import os
import re
import pickle
import time
import pdfplumber
import docx
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer

# Ensure NLTK data is ready
def check_nltk():
    for resource in ['stopwords', 'punkt', 'wordnet', 'omw-1.4']:
        try:
            if resource == 'stopwords':
                nltk.data.find('corpora/stopwords')
            elif resource == 'punkt':
                nltk.data.find('tokenizers/punkt')
            elif resource == 'wordnet':
                nltk.data.find('corpora/wordnet')
            elif resource == 'omw-1.4':
                nltk.data.find('corpora/omw-1.4')
        except LookupError:
            nltk.download(resource, quiet=True)

# Preprocessing
def preprocess_text(text):
    if not isinstance(text, str):
        return ""
    text = text.lower()
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    check_nltk()
    tokens = word_tokenize(text)
    stop_words = set(stopwords.words('english'))
    tokens = [t for t in tokens if t not in stop_words]
    lemmatizer = WordNetLemmatizer()
    tokens = [lemmatizer.lemmatize(t) for t in tokens]
    return " ".join(tokens)

# Extract Text from PDF
def extract_text_from_pdf(filepath):
    text_list = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_list.append(text)
        extracted = "\n".join(text_list).strip()
        if not extracted:
            raise ValueError("No extractable text found in PDF. The document may be scanned or empty.")
        return extracted
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")

# Extract Text from DOCX
def extract_text_from_docx(filepath):
    try:
        doc = docx.Document(filepath)
        text_list = []
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_list.append(para.text)
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_list.append(cell.text)
        extracted = "\n".join(text_list).strip()
        if not extracted:
            raise ValueError("No text found in DOCX file.")
        return extracted
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")

# Extract Text from TXT
def extract_text_from_txt(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read().strip()
        if not text:
            raise ValueError("The text file is empty.")
        return text
    except Exception as e:
        raise ValueError(f"Failed to read TXT: {str(e)}")

# Main Document Text Extractor
def extract_text(filepath):
    _, ext = os.path.splitext(filepath.lower())
    if ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext == '.docx':
        return extract_text_from_docx(filepath)
    elif ext == '.txt':
        return extract_text_from_txt(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

# Predictor Class
class DocumentPredictor:
    def __init__(self, model_path='model.pkl', vectorizer_path='vectorizer.pkl'):
        self.model_path = model_path
        self.vectorizer_path = vectorizer_path
        self.model = None
        self.vectorizer = None
        self.load_models()

    def load_models(self):
        if not os.path.exists(self.model_path) or not os.path.exists(self.vectorizer_path):
            raise FileNotFoundError("Model files (model.pkl, vectorizer.pkl) not found. Please run train.py first.")
        
        with open(self.model_path, 'rb') as f:
            self.model = pickle.load(f)
        with open(self.vectorizer_path, 'rb') as f:
            self.vectorizer = pickle.load(f)

    def predict(self, filepath):
        start_time = time.time()
        
        # Extract Text
        text = extract_text(filepath)
        
        # Preprocess
        preprocessed = preprocess_text(text)
        if not preprocessed.strip():
            raise ValueError("Document does not contain enough clean English words for classification.")
            
        # Vectorize
        X = self.vectorizer.transform([preprocessed])
        
        # Predict Class and Probability
        prediction = self.model.predict(X)[0]
        
        # Confidence estimation
        try:
            probabilities = self.model.predict_proba(X)[0]
            class_idx = list(self.model.classes_).index(prediction)
            confidence = probabilities[class_idx]
        except AttributeError:
            # Fallback if model doesn't support predict_proba
            confidence = 1.0

        processing_time = time.time() - start_time
        
        # Get a snippet of the raw text for preview (max 200 chars)
        preview_text = text[:300] + "..." if len(text) > 300 else text

        # Log confidence score internally for debugging
        print(f"[DEBUG] File: {os.path.basename(filepath)} | Predicted: {prediction} | Confidence: {round(confidence * 100, 2)}%")

        return {
            "filename": os.path.basename(filepath),
            "prediction": prediction,
            "confidence": round(confidence * 100, 2),
            "processing_time": round(processing_time, 4),
            "preview": preview_text
        }

if __name__ == "__main__":
    # Test execution
    try:
        predictor = DocumentPredictor()
        print("Models loaded successfully.")
    except Exception as e:
        print(f"Error checking models: {e}")
